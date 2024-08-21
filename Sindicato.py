import streamlit as st
from docx import Document
from docx.shared import Inches
from datetime import datetime
import io

# Função para salvar os dados em um documento Word e retornar o buffer do arquivo
def save_data(tipo, nome, email, número de telefone mensagem, numero_queixa_sugestao, rodape_foto):
    doc = Document()
    
    # Adiciona título
    doc.add_heading(f'{tipo} Nº {numero_queixa_sugestao}', 0)
    
    # Adiciona detalhes da queixa, sugestão ou reconhecimento
    doc.add_paragraph(f'Data: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Nome: {nome if nome else "Anónimo"}')
    doc.add_paragraph(f'Email: {email if email else "Não fornecido"}')
    doc.add_paragraph(f'Mensagem:')
    doc.add_paragraph(mensagem)
    
    # Adiciona rodapé com a imagem fornecida pelo usuário
    if rodape_foto:
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        run = footer_paragraph.add_run()
        run.add_picture(rodape_foto, width=Inches(1.25))
    
    # Salva o documento em um buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
st.set_page_config(
    page_title="Sindicato"
)

# Título do App
st.title("Página do Sindicato - Queixas, Sugestões e Reconhecimentos")

# Menu de navegação
st.sidebar.title("Menu")
menu = st.sidebar.radio("Escolha uma opção", ['Início', 'Enviar Queixa', 'Enviar Sugestão', 'Enviar Carta de Reconhecimento'])

# Variável para armazenar o número da queixa, sugestão ou reconhecimento
numero_queixa_sugestao = 1

# Página inicial
if menu == 'Início':
    st.header("Bem-vindo à Página do Sindicato!")
    st.write("Aqui podes enviar as tuas queixas, sugestões e cartas de reconhecimento de forma anónima ou com o teunome e email.")

# Página de Enviar Queixa
elif menu == 'Enviar Queixa':
    st.header("Enviar Queixa")
    nome = st.text_input("Nome (opcional)")
    email = st.text_input("Email (opcional)")
    mensagem = st.text_area("Descreva sua queixa")
    rodape_foto = st.file_uploader("Escolha uma imagem para o rodapé", type=["jpg", "png"])
    if st.button("Enviar"):
        if rodape_foto:
            doc_buffer = save_data('Queixa', nome, email, mensagem, numero_queixa_sugestao, rodape_foto)
            st.success(f"Sua queixa Nº {numero_queixa_sugestao} foi enviada com sucesso!")
            st.download_button(
                label="Baixar Documento",
                data=doc_buffer,
                file_name=f'Queixa_{numero_queixa_sugestao}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            numero_queixa_sugestao += 1
        else:
            st.warning("Por favor, adicione uma imagem para o rodapé.")

# Página de Enviar Sugestão
elif menu == 'Enviar Sugestão':
    st.header("Enviar Sugestão")
    nome = st.text_input("Nome (opcional)")
    email = st.text_input("Email (opcional)")
    mensagem = st.text_area("Descreva sua sugestão")
    rodape_foto = st.file_uploader("Escolha uma imagem para o rodapé", type=["jpg", "png"])
    if st.button("Enviar"):
        if rodape_foto:
            doc_buffer = save_data('Sugestão', nome, email, mensagem, numero_queixa_sugestao, rodape_foto)
            st.success(f"Sua sugestão Nº {numero_queixa_sugestao} foi enviada com sucesso!")
            st.download_button(
                label="Baixar Documento",
                data=doc_buffer,
                file_name=f'Sugestão_{numero_queixa_sugestao}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            numero_queixa_sugestao += 1
        else:
            st.warning("Por favor, adicione uma imagem para o rodapé.")

# Página de Enviar Carta de Reconhecimento
elif menu == 'Enviar Carta de Reconhecimento':
    st.header("Enviar Carta de Reconhecimento")
    nome = st.text_input("Nome (opcional)")
    email = st.text_input("Email (opcional)")
    mensagem = st.text_area("Descreva o reconhecimento")
    rodape_foto = st.file_uploader("Escolha uma imagem para o rodapé", type=["jpg", "png"])
    if st.button("Enviar"):
        if rodape_foto:
            doc_buffer = save_data('Reconhecimento', nome, email, mensagem, numero_queixa_sugestao, rodape_foto)
            st.success(f"Sua carta de reconhecimento Nº {numero_queixa_sugestao} foi enviada com sucesso!")
            st.download_button(
                label="Baixar Documento",
                data=doc_buffer,
                file_name=f'Reconhecimento_{numero_queixa_sugestao}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            numero_queixa_sugestao += 1
        else:
            st.warning("Por favor, adicione uma imagem para o rodapé.")
